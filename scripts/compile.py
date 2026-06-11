#!/usr/bin/env python3
"""Build Owen_Biblical_Theology_COMPLETE.docx and .json from translation/*.md.
Run from the repository root:  python3 scripts/compile.py
The numbered markdown files in translation/ are the source of truth."""
import glob, json, re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

META = {
  "title": "Theologoumena Pantodapa (Biblical Theology)",
  "subtitle": "The Nature, Origin, Progress, and Study of True Theology, in Six Books",
  "author": "John Owen",
  "source": "Translated into English from the Latin text of William H. Goold, D.D. (Edinburgh: T. & T. Clark, 1862)",
  "notes": "Printed page numbers of the 1862 edition appear as page blocks and inline [ p. N ] markers. Sections carry Owen's Roman numerals.",
}

# ---- read markdown into a flat stream of (kind, text) ----
stream = []
for f in sorted(glob.glob('translation/[0-9]*.md')):
    buf = []
    def flush():
        t = ' '.join(buf).strip()
        if t: stream.append(('p', t))
        buf.clear()
    for ln in open(f, encoding='utf-8').read().split('\n'):
        s = ln.rstrip()
        if s.startswith('## '): flush(); stream.append(('h2', s[3:].strip()))
        elif s.startswith('# '): flush(); stream.append(('h1', s[2:].strip()))
        elif not s.strip(): flush()
        else: buf.append(s.strip())
    flush()

# ---- build docx ----
out = Document()
def center(text, size, bold=False, italic=False):
    p = out.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = 'Georgia'
center(META['title'].split(' (')[0], 28, bold=True)
center('Biblical Theology', 18, italic=True)
center(META['subtitle'], 14, italic=True)
center('', 12); center(META['author'], 16, bold=True)
center(META['source'], 11, italic=True)
center('Printed page numbers of the 1862 edition are retained in brackets, e.g. [ p. 156 ]', 10, italic=True)
def add_heading(text, level):
    h = out.add_heading(text, level=level)
    for r in h.runs: r.font.name = 'Georgia'
for kind, t in stream:
    if kind == 'h1': add_heading(t, 1)
    elif kind == 'h2': add_heading(t, 2)
    else: out.add_paragraph(t)
out.save('Owen_Biblical_Theology_COMPLETE.docx')

# ---- build json ----
data = dict(META); data['front_matter'] = []; data['books'] = []
cur_book = None; cur_chap = None
page_re = re.compile(r'^\[ p\. (\d+)')
sec_re = re.compile(r'^([IVXL]+)\.\s')
base_re = re.compile(r'^(Chapter [IVXL]+)')
def new_chapter(heading):
    global cur_chap
    cur_chap = {"heading": heading, "blocks": []}
    (cur_book["chapters"] if cur_book else data["front_matter"]).append(cur_chap)
for kind, t in stream:
    if kind == 'h1':
        if t.startswith('BOOK '):
            m = re.match(r'BOOK ([IVX]+)\.\s*(.*)', t)
            cur_book = {"book": m.group(1), "title": m.group(2), "chapters": []}
            data["books"].append(cur_book); cur_chap = None
        else:
            new_chapter(re.sub(r'^Book [IVX]+\s*[—-]\s*', '', t).strip())
    elif kind == 'h2':
        if cur_chap is None: new_chapter("")
        pm = page_re.match(t)
        cur_chap["blocks"].append({"type": "page", "page": int(pm.group(1))} if pm else {"type": "subheading", "text": t})
    else:
        if cur_chap is None: new_chapter("")
        blk = {"type": "paragraph", "text": t}
        sm = sec_re.match(t)
        if sm: blk["section"] = sm.group(1)
        elif t.startswith('Chapter argument:'): blk["type"] = "argument"
        cur_chap["blocks"].append(blk)
def norm(h):
    h = re.sub(r'\((continued|concluded)\)', '', h)
    return h.replace('Digression:', 'Digression').strip(' .:').lower()
for b in data["books"]:
    merged = []
    for c in b["chapters"]:
        c["heading"] = re.sub(r'\s*\((continued|concluded)\)', '', c["heading"]).strip()
        if merged and norm(c["heading"]) == norm(merged[-1]["heading"]) and norm(c["heading"]):
            merged[-1]["blocks"].extend(c["blocks"]); continue
        merged.append(c)
    pre = [c for c in merged if not c["heading"]]
    merged = [c for c in merged if c["heading"]]
    if pre and merged:
        lead, rest = [], merged[0]["blocks"]
        if rest and rest[0].get("type") == "page": lead, rest = [rest[0]], rest[1:]
        merged[0]["blocks"] = lead + [blk for c in pre for blk in c["blocks"]] + rest
    b["chapters"] = merged
json.dump(data, open('Owen_Biblical_Theology_COMPLETE.json','w',encoding='utf-8'), ensure_ascii=False, indent=1)

words = sum(len(p.text.split()) for p in out.paragraphs)
print('docx words:', words, '| books:', len(data['books']), '| chapters:', sum(len(b["chapters"]) for b in data["books"]))
